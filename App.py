import streamlit as st
import openai
from docx import Document
from io import BytesIO
import base64
import tempfile
import os

# ========== CONFIGURATION ==========
openai.api_key = st.secrets["openai_api_key"]  # Or set directly: openai.api_key = "your-api-key"

# ========== HELPER FUNCTIONS ==========
def generate_points(jd_text, client, duration):
    prompt = f"""
    Based on the following Job Description (JD), generate 3-5 resume bullet points for a candidate who worked at {client} from {duration}. 
    Keep them tailored, keyword-rich, and formatted for a professional resume.

    JD:
    {jd_text}
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

def insert_points_to_resume(docx_file, client, generated_points):
    doc = Document(docx_file)
    found = False
    for para in doc.paragraphs:
        if client.lower() in para.text.lower():
            found = True
            index = doc.paragraphs.index(para) + 1
            for point in generated_points.split("\n"):
                if point.strip():
                    doc.paragraphs.insert(index, doc.add_paragraph(f"• {point.strip().lstrip('• ')}"))
                    index += 1
            break
    return doc, found

def convert_to_pdf(doc):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        pdf_path = tmp_docx.name.replace(".docx", ".pdf")
        os.system(f"libreoffice --headless --convert-to pdf {tmp_docx.name} --outdir {os.path.dirname(tmp_docx.name)}")
        with open(pdf_path, "rb") as pdf_file:
            pdf_bytes = pdf_file.read()
        os.remove(tmp_docx.name)
        os.remove(pdf_path)
        return pdf_bytes

# ========== STREAMLIT UI ==========
st.set_page_config(page_title="AI Resume Updater", layout="centered")
st.title("📄 AI Resume Updater")

st.markdown("Upload your resume and job description, and get a custom-tailored resume version!")

jd_text = st.text_area("Paste Job Description (JD)")
resume_file = st.file_uploader("Upload your current resume (.docx)", type=["docx"])

client = st.text_input("Client Name (e.g., Cisco)")
duration = st.text_input("Duration (e.g., Aug 2021 – Sep 2023)")

if st.button("✨ Generate Updated Resume"):
    if not all([jd_text, resume_file, client, duration]):
        st.warning("Please provide all inputs: JD, resume, client, and duration.")
    else:
        with st.spinner("Generating resume points and updating your resume..."):
            try:
                # Step 1: Generate points
                points = generate_points(jd_text, client, duration)

                # Step 2: Update resume
                updated_doc, found = insert_points_to_resume(resume_file, client, points)

                if not found:
                    st.warning(f"Could not find the section for '{client}' in your resume. Please check formatting.")
                else:
                    # Step 3: Export as DOCX
                    docx_io = BytesIO()
                    updated_doc.save(docx_io)
                    docx_io.seek(0)

                    # Step 4: Export as PDF
                    pdf_bytes = convert_to_pdf(updated_doc)

                    st.success("✅ Resume updated!")
                    st.download_button("📄 Download DOCX", docx_io, file_name="updated_resume.docx")
                    st.download_button("📄 Download PDF", pdf_bytes, file_name="updated_resume.pdf")

            except Exception as e:
                st.error(f"Something went wrong: {e}")

st.markdown("---")
st.caption("Built by Imran with ❤️ using GPT-4")
